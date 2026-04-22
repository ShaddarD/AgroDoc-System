import os
import datetime

from django.conf import settings


class DocumentGenerator:
    """Генератор документов по заявке"""

    def __init__(self):
        self.templates_dir = os.path.join(settings.BASE_DIR, 'templates', 'docs')
        self.output_dir = os.path.join(settings.MEDIA_ROOT, 'generated_docs')

    def _sender_name(self, app):
        # ⚠️ TEMP: sender_ru model not defined; using applicant_counterparty as fallback
        counterparty = getattr(app, 'applicant_counterparty', None)
        return counterparty.name_ru if counterparty else (getattr(app, 'sender_en_manual', '') or '')

    def _sender_inn(self, app):
        # ⚠️ TEMP: sender_ru model not defined; using applicant_counterparty
        counterparty = getattr(app, 'applicant_counterparty', None)
        return counterparty.inn if counterparty else ''

    def _sender_kpp(self, app):
        # ⚠️ TEMP: sender_ru model not defined; using applicant_counterparty
        counterparty = getattr(app, 'applicant_counterparty', None)
        return counterparty.kpp if counterparty else ''

    def _sender_ogrn(self, app):
        # ⚠️ TEMP: sender_ru model not defined; using applicant_counterparty
        counterparty = getattr(app, 'applicant_counterparty', None)
        return counterparty.ogrn if counterparty else ''

    def _receiver_name(self, app):
        # ⚠️ TEMP: receiver model not defined; returns empty string
        # TODO: implement after Receiver/Sender models are added (task: #XXX)
        return ''

    def _receiver_address(self, app):
        # ⚠️ TEMP: receiver model not defined; returns empty string
        # TODO: implement after Receiver/Sender models are added (task: #XXX)
        return ''

    def _product_name_ru(self, app):
        product = getattr(app, 'product', None)
        return product.name_ru if product else ''

    def _product_botanical(self, app):
        product = getattr(app, 'product', None)
        return getattr(product, 'botanical_name_latin', '') if product else ''

    def _packing_type(self, app):
        # ⚠️ TEMP: packing_type model not defined; returns empty string
        # TODO: implement after PackingType model is added (task: #XXX)
        return ''

    def _sampling_place(self, app):
        # ⚠️ TEMP: sampling_place model not defined; returns empty string
        # TODO: implement after SamplingPlace model is added (task: #XXX)
        return ''

    def _containers_str(self, app):
        # ⚠️ TEMP: containers relation not defined; returns empty string
        # TODO: implement via dedicated containers table (task: #XXX)
        return ''

    def _output_path(self, app, filename):
        date_str = app.created_at.strftime('%Y/%m/%d')
        path = os.path.join(self.output_dir, date_str, filename)
        os.makedirs(os.path.dirname(path), exist_ok=True)
        return path

    def generate_all(self, application):
        files = []
        cokz = self.generate_cokz(application)
        if cokz:
            files.append(cokz)
        files.extend(self.generate_fito(application))
        return files

    def generate_cokz(self, application):
        from docx import Document

        template_path = os.path.join(self.templates_dir, 'cokz_template.docx')
        if os.path.exists(template_path):
            doc = Document(template_path)
        else:
            doc = Document()
            doc.add_heading('ЗАЯВКА В ЦОКЗ', 0)
            doc.add_paragraph(f'Номер заявки: {application.application_number}')
            doc.add_paragraph(f'Отправитель: {self._sender_name(application)}')
            doc.add_paragraph(f'ИНН: {self._sender_inn(application)}')
            doc.add_paragraph(f'КПП: {self._sender_kpp(application)}')
            doc.add_paragraph(f'ОГРН: {self._sender_ogrn(application)}')
            doc.add_paragraph(f'Получатель: {self._receiver_name(application)}')
            doc.add_paragraph(f'Адрес получателя: {self._receiver_address(application)}')
            doc.add_paragraph(f'Продукция (рус): {self._product_name_ru(application)}')
            doc.add_paragraph(f'Продукция (eng): {application.product_name_en_manual}')
            doc.add_paragraph(f'Ботаническое название: {self._product_botanical(application)}')
            doc.add_paragraph(f'Вес: {application.weight_mt} MT')
            doc.add_paragraph(f'Тип упаковки: {self._packing_type(application)}')
            doc.add_paragraph(f'Контейнеры: {self._containers_str(application)}')
            doc.add_paragraph(f'Порт выгрузки (рус): {application.discharge_port_ru_manual}')
            doc.add_paragraph(f'Порт выгрузки (eng): {application.discharge_port_en_manual}')
            doc.add_paragraph(f'Место отбора: {self._sampling_place(application)}')
            doc.add_paragraph(f'Дата инспекции: {application.planned_inspection_date}')

        stamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"cokz_{application.application_number}_{stamp}.docx"
        path = self._output_path(application, filename)
        doc.save(path)
        return {'name': filename, 'path': path, 'type': 'cokz'}

    def generate_fito(self, application):
        from openpyxl import Workbook

        files = []

        wb1 = Workbook()
        ws1 = wb1.active
        ws1.title = 'Фито (1)'
        rows1 = [
            ['Заявление на выдачу фитосанитарного сертификата'],
            [],
            ['Отправитель:', self._sender_name(application)],
            ['ИНН:', self._sender_inn(application)],
            ['КПП:', self._sender_kpp(application)],
            ['ОГРН:', self._sender_ogrn(application)],
            ['Отправитель (eng):', application.sender_en_manual],
            [],
            ['Получатель:', self._receiver_name(application)],
            ['Адрес:', self._receiver_address(application)],
            [],
            ['Продукция (рус):', self._product_name_ru(application)],
            ['Продукция (eng):', application.product_name_en_manual],
            ['Ботаническое название:', self._product_botanical(application)],
            ['Вес:', f'{application.weight_mt} MT'],
            ['Год урожая:', application.harvest_year],
            ['Дата выработки:', str(application.manufacture_date or '')],
            ['Тип упаковки:', self._packing_type(application)],
            ['Контейнеры:', self._containers_str(application)],
            ['Порт выгрузки (рус):', application.discharge_port_ru_manual],
            ['Порт выгрузки (eng):', application.discharge_port_en_manual],
            ['Место отбора проб:', self._sampling_place(application)],
            ['Дата инспекции (план):', str(application.planned_inspection_date or '')],
        ]
        for row in rows1:
            ws1.append(row)

        filename1 = f"fito1_{application.application_number}.xlsx"
        path1 = self._output_path(application, filename1)
        wb1.save(path1)
        files.append({'name': filename1, 'path': path1, 'type': 'fito1'})

        wb2 = Workbook()
        ws2 = wb2.active
        ws2.title = 'Фито (2)'
        ws2['A1'] = 'Заявление на отбор проб и установление карантинного фитосанитарного состояния'
        ws2['A3'] = f'Отправитель: {self._sender_name(application)}'
        ws2['A4'] = f'ИНН: {self._sender_inn(application)}'
        ws2['A5'] = f'Получатель: {self._receiver_name(application)}'
        ws2['A6'] = f'Продукция: {self._product_name_ru(application)}'
        ws2['A7'] = f'Вес: {application.weight_mt} MT'
        ws2['A8'] = f'Место отбора: {self._sampling_place(application)}'
        ws2['A9'] = f'Контейнеры: {self._containers_str(application)}'

        filename2 = f"fito2_{application.application_number}.xlsx"
        path2 = self._output_path(application, filename2)
        wb2.save(path2)
        files.append({'name': filename2, 'path': path2, 'type': 'fito2'})

        wb3 = Workbook()
        ws3 = wb3.active
        ws3.title = 'Акт досмотра'
        ws3['A1'] = 'АКТ ДОСМОТРА ПОДКАРАНТИННЫХ МАТЕРИАЛОВ'
        ws3['A3'] = f'Продукция: {self._product_name_ru(application)}'
        ws3['A4'] = f'Ботаническое название: {self._product_botanical(application)}'
        ws3['A5'] = f'Контейнеры: {self._containers_str(application)}'
        ws3['A6'] = f'Отправитель: {self._sender_name(application)}'
        ws3['A7'] = f'Получатель: {self._receiver_name(application)}'
        ws3['A8'] = f'Вес: {application.weight_mt} MT'
        ws3['A9'] = f'Место отбора: {self._sampling_place(application)}'
        ws3['A10'] = f'Номер контракта: {application.contract_number_manual}'
        ws3['A11'] = f'Дата контракта: {application.contract_date_manual}'

        filename3 = f"act_{application.application_number}.xlsx"
        path3 = self._output_path(application, filename3)
        wb3.save(path3)
        files.append({'name': filename3, 'path': path3, 'type': 'act'})

        return files
